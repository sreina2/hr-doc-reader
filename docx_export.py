from io import BytesIO

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from resume_render import (
    HEADER_MODE_EVERY_PAGE,
    HEADER_MODE_FIRST_PAGE,
    HEADER_MODE_NONE,
    MAVEN_TEXT_RGB,
    OFFICIAL_HEADER_PNG,
    split_trailing_date,
)

_BOLD_FLAG = 1 << 4


def _add_picture_to(header, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = alignment
    run = paragraph.add_run()
    run.add_picture(BytesIO(OFFICIAL_HEADER_PNG), width=Inches(6.5))


def _apply_header_mode(document: Document, header_mode: str) -> None:
    """Places the fixed official Maven header image according to header_mode - never
    extracted from a source document, always this one stored asset. "First page only"
    uses Word's native different-first-page-header mechanism so later pages get no
    image at all, rather than relying on a manual per-page insertion."""
    if header_mode == HEADER_MODE_NONE or not OFFICIAL_HEADER_PNG:
        return

    section = document.sections[0]
    if header_mode == HEADER_MODE_FIRST_PAGE:
        section.different_first_page_header_footer = True
        _add_picture_to(section.first_page_header)
        # section.header applies to all pages after the first - leave it empty.
    else:
        _add_picture_to(section.header)


def build_docx_from_sections(
    sections: list, approximate_layout: bool = False, header_mode: str = HEADER_MODE_EVERY_PAGE
) -> bytes:
    """Builds a DOCX from structured {title, lines} sections - used for every level
    that already produces this structure (Level 3 always, Level 1/2 on non-PDF
    sources, both Contract/Letter levels)."""
    maven_styling = header_mode != HEADER_MODE_NONE
    document = Document()
    _apply_header_mode(document, header_mode)

    for section in sections:
        heading = document.add_heading(section["title"], level=2)
        if maven_styling:
            for run in heading.runs:
                run.font.color.rgb = RGBColor(*MAVEN_TEXT_RGB)

        for line in section["lines"]:
            split = split_trailing_date(line) if approximate_layout else None
            if split:
                left, right = split
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(6.3), WD_TAB_ALIGNMENT.RIGHT
                )
                paragraph.add_run(left)
                paragraph.add_run("\t" + right)
            elif line.startswith("- "):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                document.add_paragraph(line)

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_docx_from_redacted_pdf(pdf_bytes: bytes, header_mode: str = HEADER_MODE_EVERY_PAGE) -> bytes:
    """Best-effort DOCX reconstruction from a redacted true-layout PDF (Level 1/2 on
    PDF sources). Word's flowed-paragraph model can't reproduce a PDF's pixel layout
    exactly, so this rebuilds paragraph structure and bold emphasis from the PDF's own
    font metadata rather than attempting a visual pixel-for-pixel match - it will read
    as a clean, properly formatted document, but line-by-line spacing and right-aligned
    dates from the original page won't carry over exactly."""
    document = Document()
    _apply_header_mode(document, header_mode)

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                text = "".join(s.get("text", "") for s in spans).strip()
                if not text:
                    continue
                is_bold = any(
                    (s.get("flags", 0) & _BOLD_FLAG) or "bold" in s.get("font", "").lower()
                    for s in spans
                )
                size = max((s.get("size", 10) for s in spans), default=10)
                paragraph = document.add_paragraph()
                run = paragraph.add_run(text)
                run.bold = bool(is_bold)
                if size >= 13:
                    run.font.size = Pt(min(size, 18))
    doc.close()

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()
